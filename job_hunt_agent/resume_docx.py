"""Deterministic, text-faithful DOCX exports for approved tailored resumes."""

from __future__ import annotations

import io
import re
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_FIXED_ARCHIVE_TIME = (1980, 1, 1, 0, 0, 0)
_FIXED_DOCUMENT_TIME = datetime(2000, 1, 1, tzinfo=timezone.utc)
_UNSAFE_FILENAME = re.compile(r"[^a-z0-9]+")
_BULLET_PREFIXES = ("- ", "• ", "● ")
_HEADING_MAX_CHARS = 80


class ResumeDocxError(ValueError):
    """The exact saved text cannot be represented safely in a DOCX file."""


@dataclass(frozen=True)
class ApprovedResumeDocx:
    """One owner-scoped, exact approved artifact prepared for download."""

    content: bytes
    filename: str
    artifact_revision_id: str
    content_hash: str


def build_resume_docx(text: str) -> bytes:
    """Return a deterministic, single-column DOCX containing only ``text`` lines."""

    try:
        document = Document()
        _configure_document(document)
        lines = text.split("\n")
        name_line_index = _name_line_index(lines)
        for index, line in enumerate(lines):
            paragraph = document.add_paragraph(
                style=_paragraph_style(line, index=index, name_line_index=name_line_index)
            )
            paragraph.add_run(line)
        raw = io.BytesIO()
        document.save(raw)
    except (TypeError, ValueError) as exc:
        raise ResumeDocxError("resume text contains unsupported document characters") from exc
    return _normalize_archive(raw.getvalue())


def safe_resume_filename(
    *,
    company_name: str,
    role_title: str,
    revision_number: int,
) -> str:
    """Build a short ASCII attachment name without trusting posting punctuation."""

    source = unicodedata.normalize("NFKD", f"{company_name} {role_title}")
    ascii_source = source.encode("ascii", "ignore").decode("ascii").casefold()
    stem = _UNSAFE_FILENAME.sub("-", ascii_source).strip("-")[:80].rstrip("-")
    if not stem:
        stem = "application"
    return f"{stem}-resume-r{revision_number}.docx"


def _configure_document(document: DocxDocument) -> None:
    properties = document.core_properties
    properties.title = ""
    properties.subject = ""
    properties.author = ""
    properties.keywords = ""
    properties.comments = ""
    properties.last_modified_by = ""
    properties.revision = 1
    properties.created = _FIXED_DOCUMENT_TIME
    properties.modified = _FIXED_DOCUMENT_TIME

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    _set_style_font(normal, size=Pt(10.5), bold=False)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 1.08

    heading = document.styles.add_style("Resume Heading", WD_STYLE_TYPE.PARAGRAPH)
    heading.base_style = normal
    _set_style_font(heading, size=Pt(11), bold=True)
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(2)
    heading.paragraph_format.line_spacing = 1.0
    heading.paragraph_format.keep_with_next = True

    name = document.styles.add_style("Resume Name", WD_STYLE_TYPE.PARAGRAPH)
    name.base_style = normal
    _set_style_font(name, size=Pt(15), bold=True)
    name.paragraph_format.space_before = Pt(0)
    name.paragraph_format.space_after = Pt(2)
    name.paragraph_format.line_spacing = 1.0
    name.paragraph_format.keep_with_next = True

    bullet = document.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = normal
    _set_style_font(bullet, size=Pt(10.5), bold=False)
    bullet.paragraph_format.left_indent = Inches(0.18)
    bullet.paragraph_format.first_line_indent = Inches(-0.18)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(1)
    bullet.paragraph_format.line_spacing = 1.08

    spacer = document.styles.add_style("Resume Spacer", WD_STYLE_TYPE.PARAGRAPH)
    spacer.base_style = normal
    _set_style_font(spacer, size=Pt(6), bold=False)
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 1.0


def _set_style_font(style: object, *, size: Pt, bold: bool) -> None:
    font = style.font  # type: ignore[attr-defined]
    font.name = "Arial"
    font.size = size
    font.bold = bold
    font.color.rgb = RGBColor(0, 0, 0)
    r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()  # type: ignore[attr-defined]
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:eastAsia"), "Arial")


def _paragraph_style(line: str, *, index: int, name_line_index: int | None) -> str:
    if not line:
        return "Resume Spacer"
    if index == name_line_index:
        return "Resume Name"
    if line.startswith(_BULLET_PREFIXES):
        return "Resume Bullet"
    if _is_heading(line):
        return "Resume Heading"
    return "Normal"


def _name_line_index(lines: list[str]) -> int | None:
    if not lines:
        return None
    start = 0
    if lines[0].strip() == "RELEVANT HIGHLIGHTS":
        separator = next((index for index in range(1, len(lines)) if not lines[index]), None)
        if separator is None:
            return None
        start = separator + 1
    return next((index for index in range(start, len(lines)) if lines[index]), None)


def _is_heading(line: str) -> bool:
    stripped = line.strip()
    letters = [character for character in stripped if character.isalpha()]
    return bool(letters) and len(stripped) <= _HEADING_MAX_CHARS and stripped == stripped.upper()


def _normalize_archive(raw: bytes) -> bytes:
    source = io.BytesIO(raw)
    result = io.BytesIO()
    with zipfile.ZipFile(source, "r") as input_archive:
        with zipfile.ZipFile(
            result,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as output_archive:
            for name in input_archive.namelist():
                info = zipfile.ZipInfo(name, date_time=_FIXED_ARCHIVE_TIME)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.create_system = 0
                info.external_attr = 0
                output_archive.writestr(
                    info,
                    input_archive.read(name),
                    compress_type=zipfile.ZIP_DEFLATED,
                    compresslevel=9,
                )
    return result.getvalue()


__all__ = [
    "DOCX_MEDIA_TYPE",
    "ApprovedResumeDocx",
    "ResumeDocxError",
    "build_resume_docx",
    "safe_resume_filename",
]
