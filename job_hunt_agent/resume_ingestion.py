"""Bounded, provider-free extraction of owner-uploaded resume files.

The original file is treated as untrusted input and is never retained by this
module.  Callers persist only the normalized extracted text and resume-backed
facts returned in :class:`ParsedResume`.
"""

from __future__ import annotations

import calendar
import io
import re
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import PurePath

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from .security import MAX_RESUME_CHARS


MAX_RESUME_FILE_BYTES = 3 * 1024 * 1024
MAX_PDF_PAGES = 10
MAX_DOCX_ARCHIVE_ENTRIES = 500
MAX_DOCX_UNCOMPRESSED_BYTES = 20 * 1024 * 1024
MAX_IMPORTED_EVIDENCE = 20
RESUME_PARSER_VERSION = "1"

_GENERIC_MEDIA_TYPES = {"", "application/octet-stream"}
_PDF_MEDIA_TYPES = {"application/pdf", *_GENERIC_MEDIA_TYPES}
_DOCX_MEDIA_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/zip",
    *_GENERIC_MEDIA_TYPES,
}
_TEXT_MEDIA_TYPES = {"text/plain", *_GENERIC_MEDIA_TYPES}
_ZIP_SIGNATURES = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")
_ZERO_WIDTH = dict.fromkeys(map(ord, "\u200b\u200c\u200d\u2060\ufeff"), None)
_BULLET_RE = re.compile(r"^\s*(?:[•●▪◦‣]|[-*])\s+")
_STANDALONE_BULLET_RE = re.compile(r"^\s*(?:[•●▪◦‣]|[-*])\s*$")
_TERMINAL_RE = re.compile(r"[.!?)](?:[\"')\]]+)?$")
_ACTION_RE = re.compile(
    r"\b(?:achieved|architected|automated|built|created|cut|delivered|deployed|"
    r"designed|developed|drove|enabled|established|grew|implemented|improved|"
    r"increased|launched|led|lifted|managed|migrated|optimized|orchestrated|"
    r"owned|processed|published|ran|reduced|resolved|scaled|shipped|trained)\b",
    re.IGNORECASE,
)
_MONTHS = {
    "jan": 1,
    "january": 1,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "march": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "jun": 6,
    "june": 6,
    "jul": 7,
    "july": 7,
    "aug": 8,
    "august": 8,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,
}
_MONTH_PATTERN = "|".join(sorted(_MONTHS, key=len, reverse=True))
_DATE_RANGE_RE = re.compile(
    rf"(?P<start_month>\b(?:{_MONTH_PATTERN})\b)?\s*"
    r"(?P<start_year>(?:19|20)\d{2})\s*"
    r"(?:-|–|—|to)\s*"
    rf"(?:(?P<present>present|current|now)|(?P<end_month>\b(?:{_MONTH_PATTERN})\b)?\s*"
    r"(?P<end_year>(?:19|20)\d{2}))",
    re.IGNORECASE,
)
_SECTION_ALIASES = {
    "profile": "profile",
    "professional profile": "profile",
    "professional summary": "profile",
    "summary": "profile",
    "about": "profile",
    "professional experience": "experience",
    "work experience": "experience",
    "work history": "experience",
    "employment history": "experience",
    "experience": "experience",
    "education": "education",
    "academic background": "education",
    "projects": "projects",
    "selected projects": "projects",
    "personal projects": "projects",
    "skills": "skills",
    "technical skills": "skills",
    "core skills": "skills",
    "technologies": "skills",
    "certificates": "certifications",
    "certifications": "certifications",
    "licenses & certifications": "certifications",
    "licenses and certifications": "certifications",
    "achievements": "achievements",
    "awards": "achievements",
    "honors & awards": "achievements",
    "honors and awards": "achievements",
}
_EVIDENCE_SECTIONS = {"experience", "projects", "achievements"}


class ResumeIngestionError(ValueError):
    """Safe, user-actionable rejection of an uploaded resume."""

    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


@dataclass(frozen=True)
class ParsedEvidence:
    statement: str
    source_excerpt: str
    skills: tuple[str, ...]


@dataclass(frozen=True)
class ParsedResume:
    content: str
    sections: tuple[str, ...]
    current_title: str | None
    current_location: str | None
    years_of_experience: float | None
    evidence: tuple[ParsedEvidence, ...]
    skills: tuple[str, ...]
    warnings: tuple[str, ...]
    media_type: str
    page_count: int | None
    parser_version: str = RESUME_PARSER_VERSION


def parse_resume_upload(
    data: bytes,
    *,
    filename: str,
    content_type: str | None,
    as_of: date | None = None,
) -> ParsedResume:
    """Validate, extract, and conservatively structure one resume upload."""

    safe_name = _validated_filename(filename)
    if not data:
        raise ResumeIngestionError("resume_empty", "This resume file is empty.")
    if len(data) > MAX_RESUME_FILE_BYTES:
        raise ResumeIngestionError(
            "resume_too_large",
            "Resume files must be 3 MB or smaller.",
        )

    extension = PurePath(safe_name).suffix.casefold()
    media_type = (content_type or "").split(";", 1)[0].strip().casefold()
    if extension == ".pdf":
        _require_media_type(media_type, _PDF_MEDIA_TYPES)
        text, page_count = _extract_pdf(data)
        canonical_media_type = "application/pdf"
    elif extension == ".docx":
        _require_media_type(media_type, _DOCX_MEDIA_TYPES)
        text = _extract_docx(data)
        page_count = None
        canonical_media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif extension == ".txt":
        _require_media_type(media_type, _TEXT_MEDIA_TYPES)
        text = _extract_text(data)
        page_count = None
        canonical_media_type = "text/plain"
    else:
        raise ResumeIngestionError(
            "resume_type_unsupported",
            "Choose a PDF, DOCX, or TXT resume.",
        )

    content = _normalize_extracted_text(text)
    if not content:
        if extension == ".pdf":
            raise ResumeIngestionError(
                "resume_pdf_needs_ocr",
                "This PDF appears to be scanned or image-only. Upload a text-based PDF or DOCX file.",
            )
        raise ResumeIngestionError(
            "resume_text_missing",
            "No readable resume text was found in this file.",
        )
    if len(content) > MAX_RESUME_CHARS:
        raise ResumeIngestionError(
            "resume_text_too_large",
            f"The extracted resume text exceeds {MAX_RESUME_CHARS:,} characters.",
        )
    if extension == ".pdf" and len(re.sub(r"\W", "", content)) < 80:
        raise ResumeIngestionError(
            "resume_pdf_needs_ocr",
            "This PDF appears to be scanned or image-only. Upload a text-based PDF or DOCX file.",
        )

    section_order, sections = _split_sections(content)
    skills = _extract_skills(sections.get("skills", []), content)
    title = _current_title(sections.get("experience", []))
    experience = _experience_years(
        sections.get("experience", []),
        as_of=as_of or date.today(),
    )
    evidence = _extract_evidence(sections, skills)
    warnings: list[str] = []
    if not section_order:
        warnings.append(
            "Section headings were not detected, so fewer profile details may have been imported."
        )
    if title is None:
        warnings.append("We could not confidently identify a current job title.")
    if experience is not None:
        warnings.append(
            "Experience length is estimated from the work dates in your resume; review it if anything is missing."
        )
    if not evidence:
        warnings.append(
            "No clear resume-backed achievement bullets were found; you can add achievements later."
        )

    return ParsedResume(
        content=content,
        sections=tuple(section_order),
        current_title=title,
        current_location=None,
        years_of_experience=experience,
        evidence=tuple(evidence),
        skills=tuple(skills),
        warnings=tuple(warnings),
        media_type=canonical_media_type,
        page_count=page_count,
    )


def _validated_filename(filename: str) -> str:
    normalized = unicodedata.normalize("NFKC", filename or "").translate(_ZERO_WIDTH)
    normalized = normalized.replace("\\", "/").rsplit("/", 1)[-1].strip()
    if not normalized or len(normalized) > 255 or "\x00" in normalized:
        raise ResumeIngestionError(
            "resume_filename_invalid",
            "The resume filename is missing or invalid.",
        )
    return normalized


def _require_media_type(value: str, allowed: set[str]) -> None:
    if value not in allowed:
        raise ResumeIngestionError(
            "resume_type_mismatch",
            "The file content type does not match its PDF, DOCX, or TXT extension.",
        )


def _extract_pdf(data: bytes) -> tuple[str, int]:
    if b"%PDF-" not in data[:1024]:
        raise ResumeIngestionError(
            "resume_pdf_invalid",
            "This file does not contain a valid PDF signature.",
        )
    try:
        # Real resumes are commonly produced by browser-based design tools
        # with recoverable spec quirks. pypdf's non-strict reader still keeps
        # the explicit size/page/encryption bounds above while accepting them.
        reader = PdfReader(io.BytesIO(data), strict=False)
        if reader.is_encrypted:
            raise ResumeIngestionError(
                "resume_pdf_encrypted",
                "Password-protected PDFs are not supported. Upload an unlocked copy.",
            )
        page_count = len(reader.pages)
        if page_count == 0:
            raise ResumeIngestionError("resume_pdf_invalid", "This PDF has no pages.")
        if page_count > MAX_PDF_PAGES:
            raise ResumeIngestionError(
                "resume_pdf_too_many_pages",
                f"Resume PDFs must have {MAX_PDF_PAGES} pages or fewer.",
            )
        pages: list[str] = []
        for page in reader.pages:
            if "/Contents" not in page:
                pages.append("")
                continue
            try:
                pages.append(page.extract_text(extraction_mode="layout") or "")
            except TypeError:
                # Compatibility fallback for a valid pypdf implementation that
                # does not expose layout extraction for a particular page.
                pages.append(page.extract_text() or "")
    except ResumeIngestionError:
        raise
    except (PdfReadError, OSError, ValueError, TypeError, KeyError) as exc:
        raise ResumeIngestionError(
            "resume_pdf_invalid",
            "This PDF is damaged or cannot be read safely.",
        ) from exc
    return "\n\n".join(pages), page_count


def _extract_docx(data: bytes) -> str:
    if not data.startswith(_ZIP_SIGNATURES):
        raise ResumeIngestionError(
            "resume_docx_invalid",
            "This file does not contain a valid DOCX document.",
        )
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            entries = archive.infolist()
            names = {entry.filename for entry in entries}
            if len(entries) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise ResumeIngestionError(
                    "resume_docx_unsafe",
                    "This DOCX contains too many embedded parts to process safely.",
                )
            if sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ResumeIngestionError(
                    "resume_docx_unsafe",
                    "This DOCX expands beyond the safe processing limit.",
                )
            if any(entry.flag_bits & 0x1 for entry in entries):
                raise ResumeIngestionError(
                    "resume_docx_encrypted",
                    "Password-protected DOCX files are not supported.",
                )
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ResumeIngestionError(
                    "resume_docx_invalid",
                    "This file is not a valid DOCX document.",
                )
            if any(name.casefold().endswith("vbaproject.bin") for name in names):
                raise ResumeIngestionError(
                    "resume_docx_unsafe",
                    "Macro-enabled Word files are not supported. Save the resume as DOCX.",
                )
    except ResumeIngestionError:
        raise
    except (zipfile.BadZipFile, OSError, ValueError) as exc:
        raise ResumeIngestionError(
            "resume_docx_invalid",
            "This DOCX is damaged or cannot be read safely.",
        ) from exc

    try:
        document = Document(io.BytesIO(data))
        lines: list[str] = []
        for section in document.sections:
            lines.extend(
                paragraph.text for paragraph in section.header.paragraphs if paragraph.text.strip()
            )
        for item in document.iter_inner_content():
            if isinstance(item, Paragraph):
                text = item.text
                if text.strip() and _paragraph_is_list(item) and not _BULLET_RE.match(text):
                    text = f"• {text}"
                lines.append(text)
            elif isinstance(item, Table):
                for row in item.rows:
                    cells: list[str] = []
                    for cell in row.cells:
                        value = " ".join(cell.text.split())
                        if value and (not cells or value != cells[-1]):
                            cells.append(value)
                    if cells:
                        lines.append(" | ".join(cells))
        for section in document.sections:
            lines.extend(
                paragraph.text for paragraph in section.footer.paragraphs if paragraph.text.strip()
            )
    except (OSError, ValueError, KeyError) as exc:
        raise ResumeIngestionError(
            "resume_docx_invalid",
            "This DOCX is damaged or cannot be read safely.",
        ) from exc
    return "\n".join(lines)


def _paragraph_is_list(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name if paragraph.style is not None else "").casefold()
    properties = paragraph._p.pPr
    return "list" in style_name or (properties is not None and properties.numPr is not None)


def _extract_text(data: bytes) -> str:
    if b"\x00" in data:
        raise ResumeIngestionError(
            "resume_text_invalid",
            "This TXT file appears to be binary rather than plain text.",
        )
    try:
        value = data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ResumeIngestionError(
            "resume_text_invalid",
            "TXT resumes must use UTF-8 text encoding.",
        ) from exc
    control_count = sum(
        1 for character in value if ord(character) < 32 and character not in "\n\r\t"
    )
    if control_count > max(2, len(value) // 100):
        raise ResumeIngestionError(
            "resume_text_invalid",
            "This TXT file contains unsupported binary control characters.",
        )
    return value


def _normalize_extracted_text(value: str) -> str:
    normalized = unicodedata.normalize("NFKC", value).translate(_ZERO_WIDTH)
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\u00ad", "").replace("\u00a0", " ")
    lines: list[str] = []
    prior_blank = True
    pending_bullet = False
    for raw in normalized.split("\n"):
        line = " ".join(raw.replace("\t", " ").split()).strip()
        line = re.sub(r"^([•●▪◦‣])(?=\S)", r"\1 ", line)
        if _STANDALONE_BULLET_RE.fullmatch(line):
            pending_bullet = True
            continue
        if line:
            if pending_bullet and not _BULLET_RE.match(line):
                line = f"• {line}"
            pending_bullet = False
            lines.append(line)
            prior_blank = False
        elif not prior_blank:
            lines.append("")
            prior_blank = True
    while lines and not lines[-1]:
        lines.pop()
    return "\n".join(lines)


def _split_sections(content: str) -> tuple[list[str], dict[str, list[str]]]:
    order: list[str] = []
    sections: dict[str, list[str]] = {"header": []}
    current = "header"
    for line in content.split("\n"):
        heading = _section_heading(line)
        if heading is not None:
            current = heading
            if heading not in order:
                order.append(heading)
            sections.setdefault(heading, [])
            continue
        sections.setdefault(current, []).append(line)
    return order, sections


def _section_heading(line: str) -> str | None:
    normalized = _BULLET_RE.sub("", line).strip().strip(":").strip()
    if not normalized or len(normalized) > 60:
        return None
    return _SECTION_ALIASES.get(normalized.casefold())


def _current_title(lines: list[str]) -> str | None:
    nonempty = [(index, line) for index, line in enumerate(lines) if line]
    for position, (index, line) in enumerate(nonempty[:20]):
        match = next(
            (candidate for candidate in _DATE_RANGE_RE.finditer(line) if candidate.group("present")),
            None,
        )
        if match is None:
            continue
        same_line = _clean_title_candidate(f"{line[:match.start()]} {line[match.end():]}")
        if same_line is not None:
            return same_line
        if position > 0 and nonempty[position - 1][0] == index - 1:
            previous = _clean_title_candidate(nonempty[position - 1][1])
            if previous is not None:
                return previous
    return None


def _clean_title_candidate(value: str) -> str | None:
    candidate = _BULLET_RE.sub("", value)
    candidate = candidate.strip(" |·•-–—,")
    candidate = " ".join(candidate.split())
    if not candidate or len(candidate) > 200 or len(candidate.split()) > 14:
        return None
    folded = candidate.casefold()
    if "@" in candidate or "http" in folded or "linkedin" in folded:
        return None
    if not any(character.isalpha() for character in candidate):
        return None
    return candidate


def _experience_years(lines: list[str], *, as_of: date) -> float | None:
    intervals: list[tuple[date, date]] = []
    for line in lines:
        for match in _DATE_RANGE_RE.finditer(line):
            start_year = int(match.group("start_year"))
            start_month = _month_number(match.group("start_month")) or 1
            start = date(start_year, start_month, 1)
            if match.group("present"):
                end = as_of
            else:
                end_year = int(match.group("end_year"))
                end_month = _month_number(match.group("end_month")) or 12
                end = date(end_year, end_month, calendar.monthrange(end_year, end_month)[1])
            if start > end or (end - start).days > 60 * 366:
                continue
            intervals.append((start, end))
    if not intervals:
        return None
    intervals.sort()
    merged: list[list[date]] = []
    for start, end in intervals:
        if not merged or start > merged[-1][1]:
            merged.append([start, end])
        elif end > merged[-1][1]:
            merged[-1][1] = end
    days = sum((end - start).days for start, end in merged)
    return round(days / 365.2425, 1)


def _month_number(value: str | None) -> int | None:
    return _MONTHS.get(value.casefold()) if value else None


def _extract_skills(lines: list[str], content: str) -> list[str]:
    skills: list[str] = []
    seen: set[str] = set()
    for line in lines:
        value = line.split(":", 1)[1] if ":" in line else line
        for raw in re.split(r"[,|;•]", value):
            skill = " ".join(raw.strip(" .:-–—").split())
            if not skill or len(skill) > 80 or len(skill.split()) > 7:
                continue
            folded = skill.casefold()
            if folded not in seen:
                seen.add(folded)
                skills.append(skill)
    known = (
        "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Go", "Rust",
        "SQL", "Node.js", "FastAPI", "Django", "Spring", "PostgreSQL", "MySQL",
        "MongoDB", "Redis", "Kafka", "Docker", "Kubernetes", "AWS", "Azure", "GCP",
        "REST APIs", "GraphQL", "OAuth 2.0", "OIDC", "SCIM", "Linux", "Git",
        "PyTorch", "TensorFlow", "SQLAlchemy",
    )
    folded_content = content.casefold()
    for skill in known:
        if skill.casefold() in folded_content and skill.casefold() not in seen:
            seen.add(skill.casefold())
            skills.append(skill)
    return skills[:80]


def _extract_evidence(
    sections: dict[str, list[str]],
    skills: list[str],
) -> list[ParsedEvidence]:
    results: list[ParsedEvidence] = []
    seen: set[str] = set()
    for section_name in ("experience", "projects", "achievements"):
        lines = sections.get(section_name, [])
        index = 0
        while index < len(lines):
            line = lines[index]
            starts_bullet = bool(_BULLET_RE.match(line))
            starts_action = bool(_ACTION_RE.search(line))
            if not line or (not starts_bullet and not starts_action):
                index += 1
                continue
            excerpt_lines = [line]
            index += 1
            while index < len(lines) and lines[index]:
                next_line = lines[index]
                if _BULLET_RE.match(next_line) or _DATE_RANGE_RE.search(next_line):
                    break
                if _TERMINAL_RE.search(excerpt_lines[-1]):
                    break
                excerpt_lines.append(next_line)
                index += 1
            excerpt = "\n".join(excerpt_lines).strip()
            normalized_statement = " ".join(
                _BULLET_RE.sub("", part).strip() for part in excerpt_lines if part.strip()
            )
            normalized_statement = " ".join(normalized_statement.split())
            if (
                len(normalized_statement) < 40
                or len(excerpt) > 1_000
                or not _ACTION_RE.search(normalized_statement)
            ):
                continue
            key = normalized_statement.casefold()
            if key in seen:
                continue
            seen.add(key)
            results.append(
                ParsedEvidence(
                    # Auto-approved evidence is deliberately byte-for-byte the
                    # canonical resume excerpt. Presentation layers collapse
                    # wrapping whitespace, while provenance stays exact.
                    statement=excerpt,
                    source_excerpt=excerpt,
                    skills=tuple(_skills_for_statement(normalized_statement, skills)),
                )
            )
            if len(results) >= MAX_IMPORTED_EVIDENCE:
                return results
    return results


def _skills_for_statement(statement: str, skills: list[str]) -> list[str]:
    folded = statement.casefold()
    matches: list[str] = []
    for skill in skills:
        token = skill.casefold()
        if token in folded:
            matches.append(skill)
        if len(matches) >= 12:
            break
    return matches


__all__ = [
    "MAX_DOCX_ARCHIVE_ENTRIES",
    "MAX_DOCX_UNCOMPRESSED_BYTES",
    "MAX_IMPORTED_EVIDENCE",
    "MAX_PDF_PAGES",
    "MAX_RESUME_FILE_BYTES",
    "RESUME_PARSER_VERSION",
    "ParsedEvidence",
    "ParsedResume",
    "ResumeIngestionError",
    "parse_resume_upload",
]
