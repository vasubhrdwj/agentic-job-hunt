"""Text fidelity and package invariants for approved resume DOCX exports."""

from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from job_hunt_agent.resume_docx import build_resume_docx, safe_resume_filename


def test_resume_docx_is_deterministic_and_contains_only_saved_lines() -> None:
    saved = (
        "RELEVANT HIGHLIGHTS\n"
        "- Shipped at-least-once delivery with retries + jitter.\n"
        "\n"
        "Vasu Bhardwaj\n"
        "Backend Engineer · AWS / Kafka / OAuth 2.0\n"
    )

    first = build_resume_docx(saved)
    second = build_resume_docx(saved)

    assert first == second
    reopened = Document(BytesIO(first))
    assert [paragraph.text for paragraph in reopened.paragraphs] == saved.split("\n")
    assert [paragraph.style.name for paragraph in reopened.paragraphs] == [
        "Resume Heading",
        "Resume Bullet",
        "Resume Spacer",
        "Resume Name",
        "Normal",
        "Resume Spacer",
    ]
    assert reopened.styles["Resume Heading"].font.bold is True
    bullet_indent = reopened.styles["Resume Bullet"].paragraph_format.first_line_indent
    assert bullet_indent is not None
    assert abs(bullet_indent - Inches(-0.18)) < 200
    assert reopened.styles["Resume Name"].font.bold is True
    assert [
        element.text
        for element in reopened.element.body.iter(qn("w:t"))
    ] == [line for line in saved.split("\n") if line]
    assert reopened.tables == []
    assert len(reopened.inline_shapes) == 0
    assert all(
        not paragraph.text
        for section in reopened.sections
        for paragraph in [*section.header.paragraphs, *section.footer.paragraphs]
    )
    assert reopened.sections[0].page_width == Inches(8.5)
    assert reopened.sections[0].page_height == Inches(11)

    with ZipFile(BytesIO(first)) as archive:
        assert all(item.date_time == (1980, 1, 1, 0, 0, 0) for item in archive.infolist())
        core = archive.read("docProps/core.xml").decode("utf-8")
        assert "python-docx" not in core
        document_xml = archive.read("word/document.xml").decode("utf-8")
        for line in saved.split("\n"):
            if line:
                assert line.replace("&", "&amp;") in document_xml
        assert "Generated" not in document_xml
        assert "Tailored resume" not in document_xml


def test_resume_filename_is_ascii_bounded_and_deterministic() -> None:
    filename = safe_resume_filename(
        company_name='Éxample / "Labs"',
        role_title="Backend Engineer: Platform & APIs",
        revision_number=12,
    )

    assert filename == "example-labs-backend-engineer-platform-apis-resume-r12.docx"
    assert filename.isascii()
    assert "/" not in filename and '"' not in filename
