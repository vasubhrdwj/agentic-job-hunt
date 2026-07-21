"""Security and extraction coverage for first-class resume uploads."""

from __future__ import annotations

from datetime import date
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.errors import LimitReachedError
from pypdf.generic import (
    ArrayObject,
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

import job_hunt_agent.resume_ingestion as resume_ingestion
from job_hunt_agent.resume_ingestion import (
    MAX_DOCX_UNCOMPRESSED_BYTES,
    MAX_PDF_CONTENT_STREAMS_PER_PAGE,
    MAX_PDF_DECODED_CONTENT_BYTES,
    MAX_PDF_ENCODED_CONTENT_BYTES,
    MAX_PDF_OBJECTS,
    MAX_PDF_PAGES,
    ResumeIngestionError,
    parse_resume_upload,
)


SAMPLE_RESUME = """Vasu Example
PROFESSIONAL SUMMARY
Backend engineer building reliable event-driven systems.

PROFESSIONAL EXPERIENCE
Software Engineer | July 2025 – Present
Example Labs | Gurugram
• Built a Kafka ingestion pipeline processing 250+ events with AWS Lambda and OAuth 2.0.
• Reduced identity onboarding time by 60% by deploying a SCIM provisioning service.

EDUCATION
B.Tech in Computer Science, 2021 – 2025

PROJECTS
Incident Commander
• Trained an incident-response model and lifted hard-task performance from 0.35 to 0.90.

SKILLS
Python, FastAPI, Kafka, Docker, AWS, OAuth 2.0, SCIM

CERTIFICATIONS
AWS Academy Graduate
"""


def test_txt_import_extracts_grounded_profile_evidence_and_sections() -> None:
    parsed = parse_resume_upload(
        SAMPLE_RESUME.encode(),
        filename="backend-resume.txt",
        content_type="text/plain; charset=utf-8",
        as_of=date(2026, 7, 20),
    )

    assert parsed.current_title == "Software Engineer"
    assert parsed.current_location is None
    assert parsed.years_of_experience == 1.1
    assert parsed.sections == (
        "profile",
        "experience",
        "education",
        "projects",
        "skills",
        "certifications",
    )
    assert len(parsed.evidence) == 3
    assert all(item.statement == item.source_excerpt for item in parsed.evidence)
    assert all(item.source_excerpt in parsed.content for item in parsed.evidence)
    assert {"Python", "Kafka", "AWS", "SCIM"}.issubset(parsed.skills)
    assert "Gurugram" not in (parsed.current_location or "")
    assert any("estimated" in warning for warning in parsed.warnings)


def test_docx_import_reads_tables_and_list_paragraphs_in_document_order() -> None:
    document = Document()
    document.add_paragraph("PROFESSIONAL EXPERIENCE")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Backend Engineer"
    table.cell(0, 1).text = "Jan 2024 – Present"
    table.cell(1, 0).text = "Example Labs"
    table.cell(1, 1).text = "Bengaluru"
    document.add_paragraph(
        "Built a reliable API platform and reduced p95 latency by 45% using Python.",
        style="List Bullet",
    )
    document.add_paragraph("SKILLS")
    document.add_paragraph("Python, PostgreSQL, Docker")
    payload = BytesIO()
    document.save(payload)

    parsed = parse_resume_upload(
        payload.getvalue(),
        filename="resume.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        as_of=date(2026, 7, 20),
    )

    assert "Backend Engineer | Jan 2024 – Present" in parsed.content
    assert parsed.current_title == "Backend Engineer"
    assert len(parsed.evidence) == 1
    assert parsed.evidence[0].statement.startswith("• Built a reliable API platform")
    assert "Python" in parsed.evidence[0].skills


def test_pdf_import_uses_layout_text_and_enforces_page_limit() -> None:
    parsed = parse_resume_upload(
        _text_pdf(SAMPLE_RESUME),
        filename="resume.pdf",
        content_type="application/octet-stream",
        as_of=date(2026, 7, 20),
    )

    assert parsed.page_count == 1
    assert parsed.current_title == "Software Engineer"
    assert "experience" in parsed.sections
    assert len(parsed.evidence) >= 2

    writer = PdfWriter()
    for _ in range(MAX_PDF_PAGES + 1):
        writer.add_blank_page(width=612, height=792)
    oversized_pages = BytesIO()
    writer.write(oversized_pages)
    with pytest.raises(ResumeIngestionError, match="pages or fewer") as exc_info:
        parse_resume_upload(
            oversized_pages.getvalue(),
            filename="long.pdf",
            content_type="application/pdf",
        )
    assert exc_info.value.code == "resume_pdf_too_many_pages"


@pytest.mark.parametrize(
    ("filename", "media_type", "payload", "code"),
    [
        ("resume.pages", "application/octet-stream", b"content", "resume_type_unsupported"),
        ("resume.pdf", "text/plain", b"%PDF-invalid", "resume_type_mismatch"),
        ("resume.pdf", "application/pdf", b"not a pdf", "resume_pdf_invalid"),
        ("resume.txt", "text/plain", b"binary\x00content", "resume_text_invalid"),
        ("resume.txt", "text/plain", b"\xff\xfe", "resume_text_invalid"),
    ],
)
def test_invalid_uploads_fail_with_specific_safe_codes(
    filename: str,
    media_type: str,
    payload: bytes,
    code: str,
) -> None:
    with pytest.raises(ResumeIngestionError) as exc_info:
        parse_resume_upload(payload, filename=filename, content_type=media_type)
    assert exc_info.value.code == code


def test_encrypted_and_image_only_pdfs_get_recovery_guidance() -> None:
    encrypted = PdfWriter()
    encrypted.add_blank_page(width=612, height=792)
    encrypted.encrypt("secret")
    encrypted_bytes = BytesIO()
    encrypted.write(encrypted_bytes)
    with pytest.raises(ResumeIngestionError, match="Password-protected") as encrypted_error:
        parse_resume_upload(
            encrypted_bytes.getvalue(),
            filename="locked.pdf",
            content_type="application/pdf",
        )
    assert encrypted_error.value.code == "resume_pdf_encrypted"

    blank = PdfWriter()
    blank.add_blank_page(width=612, height=792)
    blank_bytes = BytesIO()
    blank.write(blank_bytes)
    with pytest.raises(ResumeIngestionError, match="scanned or image-only") as blank_error:
        parse_resume_upload(
            blank_bytes.getvalue(),
            filename="scan.pdf",
            content_type="application/pdf",
        )
    assert blank_error.value.code == "resume_pdf_needs_ocr"


def test_docx_zip_bomb_metadata_is_rejected_before_document_parsing() -> None:
    payload = BytesIO()
    with ZipFile(payload, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("word/document.xml", b" " * (MAX_DOCX_UNCOMPRESSED_BYTES + 1))

    with pytest.raises(ResumeIngestionError, match="safe processing limit") as exc_info:
        parse_resume_upload(
            payload.getvalue(),
            filename="bomb.docx",
            content_type="application/zip",
        )
    assert exc_info.value.code == "resume_docx_unsafe"


def test_malformed_docx_xml_is_a_safe_invalid_upload() -> None:
    document = Document()
    document.add_paragraph("Software Engineer")
    valid = BytesIO()
    document.save(valid)
    malformed = BytesIO()
    with ZipFile(BytesIO(valid.getvalue())) as source, ZipFile(
        malformed,
        "w",
        compression=ZIP_DEFLATED,
    ) as target:
        for entry in source.infolist():
            payload = source.read(entry.filename)
            if entry.filename == "word/document.xml":
                payload = b"<w:document xmlns:w='invalid'><w:body><w:p>"
            target.writestr(entry, payload)

    with pytest.raises(ResumeIngestionError) as exc_info:
        parse_resume_upload(
            malformed.getvalue(),
            filename="malformed.docx",
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )

    assert exc_info.value.code == "resume_docx_invalid"
    assert "damaged or cannot be read safely" in str(exc_info.value)


def test_unicode_is_normalized_without_importing_an_office_as_home() -> None:
    payload = SAMPLE_RESUME.replace("Vasu Example", "Vasu\u200b Example").replace(
        "Kafka", "Ｋａｆｋａ", 1
    )
    parsed = parse_resume_upload(
        payload.encode(),
        filename="resume.txt",
        content_type=None,
        as_of=date(2026, 7, 20),
    )

    assert "\u200b" not in parsed.content
    assert "Kafka" in parsed.content
    assert parsed.current_location is None


def test_known_skills_do_not_match_inside_other_words_or_technologies() -> None:
    parsed = parse_resume_upload(
        (
            "SKILLS\n"
            "JavaScript, Data Structures & Algorithms, MongoDB, SQLAlchemy, C++\n"
        ).encode(),
        filename="resume.txt",
        content_type="text/plain",
    )

    assert "JavaScript" in parsed.skills
    assert "MongoDB" in parsed.skills
    assert "SQLAlchemy" in parsed.skills
    assert "C++" in parsed.skills
    assert "Java" not in parsed.skills
    assert "Go" not in parsed.skills
    assert "SQL" not in parsed.skills


def test_ambiguous_technology_names_require_explicit_skill_listing() -> None:
    prose = parse_resume_upload(
        (
            "PROFESSIONAL SUMMARY\n"
            "Led go-to-market planning for the Spring 2026 product launch.\n"
            "SKILLS\nPython, Kafka\n"
        ).encode(),
        filename="resume.txt",
        content_type="text/plain",
    )
    assert "Go" not in prose.skills
    assert "Spring" not in prose.skills

    explicit = parse_resume_upload(
        "SKILLS\nGo, Spring, Python\n".encode(),
        filename="resume.txt",
        content_type="text/plain",
    )
    assert {"Go", "Spring", "Python"}.issubset(explicit.skills)


def test_pdf_layout_columns_remain_skill_boundaries() -> None:
    parsed = parse_resume_upload(
        _text_pdf(
            "PROFESSIONAL SUMMARY\n"
            "Backend engineer building reliable production services and event-driven systems.\n"
            "SKILLS\n"
            "Languages                                      Infrastructure\n"
            "Python, JavaScript, SQL                         PostgreSQL, Kafka, AWS\n"
            "Backend & Security:                            Core CS:\n"
            "FastAPI, OAuth 2.0                              Data Structures & Algorithms\n"
        ),
        filename="column-resume.pdf",
        content_type="application/pdf",
    )

    assert {"Python", "JavaScript", "SQL", "PostgreSQL", "Kafka", "AWS"}.issubset(
        parsed.skills
    )
    assert "SQL PostgreSQL" not in parsed.skills
    assert "Core CS" not in parsed.skills


@pytest.mark.parametrize(
    ("experience_header", "expected_title"),
    [
        (
            "Software Engineer\nAcme Corp | Jan 2024 – Present",
            "Software Engineer",
        ),
        (
            "Software Engineer\nAcme Corp\nJan 2024 – Present",
            "Software Engineer",
        ),
        (
            "Acme Corp | Software Engineer | Jan 2024 – Present",
            "Software Engineer",
        ),
        (
            "Software Engineer | Acme Corp | Jan 2024 – Present",
            "Software Engineer",
        ),
        ("Acme Corp | Jan 2024 – Present", None),
        ("Acme Engineering | Jan 2024 – Present", None),
        ("Acme Corp | Platform | Jan 2024 – Present", None),
    ],
)
def test_current_title_prefers_clear_roles_and_fails_closed_on_companies(
    experience_header: str,
    expected_title: str | None,
) -> None:
    parsed = parse_resume_upload(
        (
            "PROFESSIONAL EXPERIENCE\n"
            f"{experience_header}\n"
            "• Built a reliable backend service for production workloads.\n"
        ).encode(),
        filename="resume.txt",
        content_type="text/plain",
        as_of=date(2026, 7, 20),
    )

    assert parsed.current_title == expected_title


def test_wrapped_parenthetical_evidence_keeps_anti_gaming_clause_only() -> None:
    first_claim = (
        "• Designed a three-task difficulty curriculum (deductive canary regression "
        "-> third-party attribution -> silent data corruption)\n"
        "with 6-component programmatic grading and structural anti-gaming guards that "
        "penalize shortcut strategies by roughly 0.4\n"
        "score points."
    )
    second_claim = (
        "• Built a separate evaluation service that reports deterministic model scores."
    )
    parsed = parse_resume_upload(
        f"PROJECTS\n{first_claim}\n{second_claim}\n".encode(),
        filename="resume.txt",
        content_type="text/plain",
    )

    assert [item.statement for item in parsed.evidence] == [first_claim, second_claim]
    assert "6-component programmatic grading" in parsed.evidence[0].statement
    assert "structural anti-gaming guards" in parsed.evidence[0].statement
    assert second_claim not in parsed.evidence[0].statement

    heading_after_parenthesis = parse_resume_upload(
        (
            "PROJECTS\n"
            "• Designed a production-ready service architecture (Python)\n"
            "Incident Commander\n"
            "• Built a separate incident evaluation service with deterministic scores.\n"
        ).encode(),
        filename="resume.txt",
        content_type="text/plain",
    )
    assert [item.statement for item in heading_after_parenthesis.evidence] == [
        "• Designed a production-ready service architecture (Python)",
        "• Built a separate incident evaluation service with deterministic scores.",
    ]


def test_pdf_object_and_content_complexity_are_bounded() -> None:
    expanded = PdfWriter()
    expanded_page = expanded.add_blank_page(width=612, height=792)
    content = DecodedStreamObject()
    content.set_data(b"q\n" * (MAX_PDF_DECODED_CONTENT_BYTES // 2 + 1))
    expanded_page[NameObject("/Contents")] = expanded._add_object(content.flate_encode())
    expanded_bytes = BytesIO()
    expanded.write(expanded_bytes)
    with pytest.raises(ResumeIngestionError) as expanded_error:
        parse_resume_upload(
            expanded_bytes.getvalue(),
            filename="expanded.pdf",
            content_type="application/pdf",
        )
    assert expanded_error.value.code == "resume_pdf_unsafe"
    assert "safe page-content limit" in str(expanded_error.value)

    object_heavy = PdfWriter()
    object_heavy.add_blank_page(width=612, height=792)
    for _ in range(MAX_PDF_OBJECTS):
        object_heavy._add_object(DictionaryObject())
    object_heavy_bytes = BytesIO()
    object_heavy.write(object_heavy_bytes)
    with pytest.raises(ResumeIngestionError) as object_error:
        parse_resume_upload(
            object_heavy_bytes.getvalue(),
            filename="object-heavy.pdf",
            content_type="application/pdf",
        )
    assert object_error.value.code == "resume_pdf_unsafe"
    assert "too many internal objects" in str(object_error.value)

    encoded = PdfWriter()
    encoded_page = encoded.add_blank_page(width=612, height=792)
    encoded_content = DecodedStreamObject()
    encoded_content.set_data(b"q\n" * (MAX_PDF_ENCODED_CONTENT_BYTES // 2 + 1))
    encoded_page[NameObject("/Contents")] = encoded._add_object(encoded_content)
    encoded_bytes = BytesIO()
    encoded.write(encoded_bytes)
    with pytest.raises(ResumeIngestionError) as encoded_error:
        parse_resume_upload(
            encoded_bytes.getvalue(),
            filename="encoded-heavy.pdf",
            content_type="application/pdf",
        )
    assert encoded_error.value.code == "resume_pdf_unsafe"
    assert "encoded page content" in str(encoded_error.value)

    fragmented = PdfWriter()
    fragmented_page = fragmented.add_blank_page(width=612, height=792)
    streams = ArrayObject()
    for _ in range(MAX_PDF_CONTENT_STREAMS_PER_PAGE + 1):
        stream = DecodedStreamObject()
        stream.set_data(b"q\n")
        streams.append(fragmented._add_object(stream))
    fragmented_page[NameObject("/Contents")] = streams
    fragmented_bytes = BytesIO()
    fragmented.write(fragmented_bytes)
    with pytest.raises(ResumeIngestionError) as fragmented_error:
        parse_resume_upload(
            fragmented_bytes.getvalue(),
            filename="fragmented.pdf",
            content_type="application/pdf",
        )
    assert fragmented_error.value.code == "resume_pdf_unsafe"
    assert "too many content streams" in str(fragmented_error.value)


def test_pdf_object_preflight_rejects_before_unbounded_index_materialization() -> None:
    class UnreadableXref:
        def values(self) -> object:
            raise AssertionError("xref must not be traversed after oversized /Size")

    declared_oversize = type(
        "DeclaredOversizeReader",
        (),
        {
            "trailer": {"/Size": MAX_PDF_OBJECTS + 1},
            "xref": UnreadableXref(),
            "xref_objStm": {},
        },
    )()
    with pytest.raises(ResumeIngestionError) as declared_error:
        resume_ingestion._require_pdf_object_bounds(declared_oversize)
    assert declared_error.value.code == "resume_pdf_unsafe"

    class TooManyIndexedObjects:
        def __len__(self) -> int:
            return MAX_PDF_OBJECTS + 1

    dishonest_size = type(
        "DishonestSizeReader",
        (),
        {
            "trailer": {"/Size": 1},
            "xref": {0: TooManyIndexedObjects()},
            "xref_objStm": {},
        },
    )()
    with pytest.raises(ResumeIngestionError) as indexed_error:
        resume_ingestion._require_pdf_object_bounds(dishonest_size)
    assert indexed_error.value.code == "resume_pdf_unsafe"


def test_pypdf_limit_errors_are_mapped_to_safe_invalid_pdf() -> None:
    class LimitedContents:
        def get_object(self) -> object:
            raise LimitReachedError("sensitive parser detail")

    class LimitedPage:
        def raw_get(self, key: str) -> LimitedContents:
            assert key == "/Contents"
            return LimitedContents()

    with pytest.raises(ResumeIngestionError) as exc_info:
        resume_ingestion._pdf_page_content_sizes(LimitedPage())

    assert exc_info.value.code == "resume_pdf_invalid"
    assert "sensitive parser detail" not in str(exc_info.value)


def _text_pdf(text: str) -> bytes:
    """Create a small Type1-font PDF fixture without another test dependency."""

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): font_ref}
            )
        }
    )
    commands = ["BT", "/F1 9 Tf", "42 755 Td", "11 TL"]
    for index, line in enumerate(text.splitlines()):
        if index:
            commands.append("T*")
        escaped = (
            line.replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
            .replace("•", "-")
            .replace("–", "-")
        )
        commands.append(f"({escaped}) Tj")
    commands.append("ET")
    stream = DecodedStreamObject()
    stream.set_data("\n".join(commands).encode("latin-1", "replace"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    result = BytesIO()
    writer.write(result)
    return result.getvalue()
